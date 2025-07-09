from langchain_community.document_loaders import PyMuPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path
from typing import List
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import os
import pandas as pd
import logging
from docx import Document
from langchain.schema import Document as LangChainDocument
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

logger = logging.getLogger(__name__)


def get_embedding_model(api_provider: str = "auto"):
    """Get appropriate embedding model based on API provider"""
    groq_key = os.getenv('GROQ_API_KEY')
    openai_key = os.getenv('OPENAI_API_KEY')
    
    if api_provider == "groq" and groq_key:
        # Use free embedding model for Groq
        try:
            from langchain_huggingface import HuggingFaceEmbeddings
            return HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        except ImportError:
            print("HuggingFace embeddings not available, falling back to OpenAI")
            return OpenAIEmbeddings(model="text-embedding-3-small")
    elif api_provider == "openai" and openai_key:
        return OpenAIEmbeddings(model="text-embedding-3-small")
    elif api_provider == "auto":
        # Auto-detect: prefer free embedding if Groq is available
        if groq_key:
            try:
                from langchain_huggingface import HuggingFaceEmbeddings
                return HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={'device': 'cpu'},
                    encode_kwargs={'normalize_embeddings': True}
                )
            except ImportError:
                print("HuggingFace embeddings not available")
                if openai_key:
                    print("Falling back to OpenAI embeddings")
                    return OpenAIEmbeddings(model="text-embedding-3-small")
                else:
                    raise ValueError("Neither HuggingFace embeddings nor OpenAI API key available. Please install langchain-huggingface or set OPENAI_API_KEY.")
        elif openai_key:
            return OpenAIEmbeddings(model="text-embedding-3-small")
        else:
            raise ValueError("No API keys found. Please set GROQ_API_KEY or OPENAI_API_KEY environment variable.")
    else:
        raise ValueError(f"Invalid API provider '{api_provider}' or missing API key")

def load_documents_from_folder(folder_path: str) -> List:
    """
    Load documents from folder supporting multiple formats:
    - PDF files
    - Word documents (.doc, .docx)
    - Excel files (.xlsx, .xls)
    - CSV files (.csv)
    - Text files (.txt)
    """
    folder = Path(folder_path)
    all_docs = []
    supported_extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt'}
    
    print(f"📁 Scanning folder: {folder}")
    
    for file in folder.iterdir():
        if file.is_file():
            if file.suffix.lower() in supported_extensions:
                print(f"📄 Processing: {file.name}")
                
                try:
                    if file.suffix.lower() == '.pdf':
                        docs = _load_pdf(file)
                    elif file.suffix.lower() in ['.docx', '.doc']:
                        docs = _load_word_document(file)
                    elif file.suffix.lower() in ['.xlsx', '.xls']:
                        docs = _load_excel_document(file)
                    elif file.suffix.lower() == '.csv':
                        docs = _load_csv_document(file)
                    elif file.suffix.lower() == '.txt':
                        docs = _load_text_document(file)
                    else:
                        continue
                        
                    all_docs.extend(docs)
                    print(f"✅ Successfully loaded {len(docs)} document chunks from {file.name}")
                    
                except Exception as e:
                    print(f"❌ Error loading {file.name}: {e}")
                    continue
            else:
                print(f"⏭️  Skipping unsupported file type: {file.name} ({file.suffix})")
    
    print(f"📊 Total documents loaded: {len(all_docs)}")
    return all_docs

def _load_pdf(file_path: Path) -> List[LangChainDocument]:
    """Load PDF files using PyMuPDFLoader"""
    loader = PyMuPDFLoader(str(file_path))
    documents = loader.load()
    
    # Add file source metadata
    for doc in documents:
        doc.metadata['source'] = str(file_path)
        doc.metadata['file_type'] = 'pdf'
        
    return documents

def _load_word_document(file_path: Path) -> List[LangChainDocument]:
    """Load Word documents (.docx, .doc)"""
    try:
        # For .docx files
        if file_path.suffix.lower() == '.docx':
            doc = Document(str(file_path))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            content = "\n\n".join(text_content)
            
        else:
            # For .doc files, use unstructured
            from unstructured.partition.doc import partition_doc
            elements = partition_doc(str(file_path))
            content = "\n\n".join([str(element) for element in elements])
        
        # Create LangChain document
        document = LangChainDocument(
            page_content=content,
            metadata={
                'source': str(file_path),
                'file_type': 'word_document'
            }
        )
        
        return [document]
        
    except Exception as e:
        print(f"Error processing Word document {file_path}: {e}")
        return []

def _load_excel_document(file_path: Path) -> List[LangChainDocument]:
    """Load Excel files (.xlsx, .xls)"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(str(file_path))
        documents = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(str(file_path), sheet_name=sheet_name)
            
            # Convert DataFrame to text representation
            text_content = []
            
            # Add sheet name as header
            text_content.append(f"=== SHEET: {sheet_name} ===")
            
            # Add column headers
            if not df.empty:
                headers = " | ".join([str(col) for col in df.columns])
                text_content.append(f"Headers: {headers}")
                text_content.append("-" * 50)
                
                # Add rows (limit to avoid too much data)
                max_rows = min(100, len(df))  # Limit to 100 rows per sheet
                for idx, row in df.head(max_rows).iterrows():
                    row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                    text_content.append(row_text)
                
                if len(df) > max_rows:
                    text_content.append(f"... and {len(df) - max_rows} more rows")
            
            content = "\n".join(text_content)
            
            # Create document for each sheet
            document = LangChainDocument(
                page_content=content,
                metadata={
                    'source': str(file_path),
                    'file_type': 'excel',
                    'sheet_name': sheet_name,
                    'total_rows': len(df),
                    'total_columns': len(df.columns)
                }
            )
            
            documents.append(document)
        
        return documents
        
    except Exception as e:
        print(f"Error processing Excel document {file_path}: {e}")
        return []

def _load_csv_document(file_path: Path) -> List[LangChainDocument]:
    """Load CSV files"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(str(file_path), encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if df is None:
            raise ValueError("Could not decode CSV file with any supported encoding")
        
        # Convert DataFrame to text representation
        text_content = []
        
        # Add filename as header
        text_content.append(f"=== CSV FILE: {file_path.name} ===")
        
        # Add column headers
        if not df.empty:
            headers = " | ".join([str(col) for col in df.columns])
            text_content.append(f"Headers: {headers}")
            text_content.append("-" * 50)
            
            # Add rows (limit to avoid too much data)
            max_rows = min(200, len(df))  # Limit to 200 rows for CSV
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                text_content.append(row_text)
            
            if len(df) > max_rows:
                text_content.append(f"... and {len(df) - max_rows} more rows")
        
        content = "\n".join(text_content)
        
        # Create document
        document = LangChainDocument(
            page_content=content,
            metadata={
                'source': str(file_path),
                'file_type': 'csv',
                'total_rows': len(df),
                'total_columns': len(df.columns)
            }
        )
        
        return [document]
        
    except Exception as e:
        print(f"Error processing CSV document {file_path}: {e}")
        return []

def _load_text_document(file_path: Path) -> List[LangChainDocument]:
    """Load plain text files (.txt)"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError("Could not decode text file with any supported encoding")
        
        # Create document
        document = LangChainDocument(
            page_content=content,
            metadata={
                'source': str(file_path),
                'file_type': 'text',
                'character_count': len(content)
            }
        )
        
        return [document]
        
    except Exception as e:
        print(f"Error processing text document {file_path}: {e}")
        return []

def chunk_documents(documents: List, chunk_size=1000, chunk_overlap=150) -> List:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(documents)


# Embed and store chunks in Chroma
def create_vector_store(chunks, persist_dir="chroma_db", api_provider="auto", force_rebuild=False):
    import shutil
    from pathlib import Path
    
    # If force rebuild is requested, delete existing vector store
    if force_rebuild and Path(persist_dir).exists():
        logger.info(f"🗑️  Deleting existing vector store at {persist_dir}")
        shutil.rmtree(persist_dir)
    
    embedding = get_embedding_model(api_provider)
    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embedding,
        persist_directory=persist_dir
    )
    vectordb.persist()
    return vectordb

# Load existing vector store (for reuse)
def load_vector_store(persist_dir="chroma_db", api_provider="auto"):
    embedding = get_embedding_model(api_provider)
    vectordb = Chroma(
        persist_directory=persist_dir,
        embedding_function=embedding
    )
    return vectordb

# Get top-k relevant chunks using RAG
def get_relevant_docs(query, vectordb, k=5):
    retriever = vectordb.as_retriever(search_type="similarity", search_kwargs={"k": k})
    return retriever.get_relevant_documents(query)
